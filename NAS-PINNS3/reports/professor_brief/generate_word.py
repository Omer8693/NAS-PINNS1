"""
Generate Word (.docx) version of the NAS-MCO-PINNs progress report.
"""
import json, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

BASE   = os.path.dirname(os.path.abspath(__file__))
ROOT   = os.path.join(BASE, "..", "..", "..")
NP3    = os.path.join(ROOT, "NAS-PINNS3")
IMG    = os.path.join(NP3,  "web", "static", "img")
V2_RES = os.path.join(NP3,  "level8_nas_mco_pinn", "results", "v2", "results_3d_v2.json")
OUT    = os.path.join(BASE,  "NAS_PINNS3_professor_brief.docx")

with open(V2_RES) as f:
    DATA = json.load(f)

DOMAINS  = ["rectangular", "cylinder", "stacked", "lshape"]
ARCHS    = ["bayesian", "nsga2", "nsga3"]
SKIPS    = [1, 2, 4, 6]
ARCH_LBL = {"bayesian": "Bayesian", "nsga2": "NSGA-II", "nsga3": "NSGA-III"}
DOM_LBL  = {"rectangular": "Rectangular", "cylinder": "Cylinder",
             "stacked": "Stacked", "lshape": "L-Shape"}
FEM_SAVE = {1: "0%", 2: "52%", 4: "71%", 6: "81%"}

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def add_run(para, text, bold=False, italic=False, size=10, color=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color.lstrip("#")))
    return run

def heading(doc, num, title, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, f"{num}.  {title}", bold=True, size=12, color="#1a237e")
    return p

def body(doc, text, size=10):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.size = Pt(size)
    return p

def formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, text, bold=False, size=10, color="#1a237e")
    return p

def fig_caption(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(10)
    add_run(p, label + " ", bold=True, size=9, color="#333333")
    add_run(p, text,        bold=False, italic=True, size=9, color="#555555")
    return p

def add_image(doc, fname, width_inches=6.2):
    path = os.path.join(IMG, fname)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
    return p


# ── Document setup ────────────────────────────────────────────────────────────
doc = Document()

# Page margins: 2.0 cm on all sides (wider content area)
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(10)


# ── Title block ───────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after  = Pt(4)
add_run(title,
    "NAS-MCO-PINNs: Neural Architecture Search with Multi-Loss Consistency\n"
    "Optimization and Skip Operator for 3D Thermal Simulation",
    bold=True, size=14, color="#111111")

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(8)
add_run(sub, "Project Progress Report  \u2014  " + date.today().strftime("%B %Y"),
        size=10, color="#555555")

# Horizontal rule (empty bordered paragraph)
hr = doc.add_paragraph()
hr.paragraph_format.space_after = Pt(8)
pPr = hr._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "6")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), "888888")
pBdr.append(bottom)
pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — Introduction
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, 1, "Introduction")
body(doc,
    "Physics-Informed Neural Networks (PINNs) offer a mesh-free approach to solving "
    "partial differential equations governing heat transfer. In a standard PINN, the "
    "network is trained on each FEM time window [t_k, t_{k+1}], so the number of PINN "
    "training runs equals the total number of FEM time steps, limiting computational savings.")
body(doc,
    "This project presents NAS-MCO-PINNs with three contributions: "
    "(i) a residual skip operator predicting s time steps ahead in a single run; "
    "(ii) Neural Architecture Search (NAS) via Bayesian Optimization, NSGA-II, and NSGA-III; "
    "and (iii) Multi-Loss Consistency Optimization (MCO) for adaptive loss weighting.")
body(doc,
    "Validated on A356 aluminum water quenching across four 3D geometries: "
    "Rectangular Prism (1.3\u00d70.6\u00d70.4 m), Cylinder (R=0.25 m, H=0.6 m), "
    "Stacked Cubes (2\u00d70.5 m), and L-Shape (0.8\u00d70.8\u00d70.4 m, numerical FD reference). "
    "Total: 48 runs (4 domains \u00d7 3 NAS optimizers \u00d7 4 skip values).")


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Skip Operator
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, 2, "The Residual Skip Operator")
body(doc,
    "Instead of one PINN per FEM step, a single network f_\u03b8 predicts s steps ahead "
    "from the current initial condition. The predicted temperature field is:")
formula(doc,
    "\u03b8\u0302(t + s\u00b7\u0394t)  =  \u03b8_IC  +  \u03c4 \u00b7 f_\u03b8(x, y, z, \u03c4, \u03b8_IC)")
body(doc,
    "where \u03c4 = t/(s\u00b7\u0394t) \u2208 [0,1] is the normalized time coordinate. This guarantees "
    "the initial condition exactly at \u03c4=0. Skipping s steps reduces FEM calls from N to "
    "ceil(N/s). Table 1 summarizes the savings assuming 21 total time steps.")

# Table 1 — FEM savings
doc.add_paragraph().paragraph_format.space_after = Pt(2)
tbl1 = doc.add_table(rows=5, cols=4)
tbl1.style = "Table Grid"
tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_data = ["Skip Value (s)", "FEM Steps Saved", "Reduction", "PINN Runs (of 21)"]
row_data1 = [
    ["s = 1  (baseline)", "0 steps", "0%",  "21 runs"],
    ["s = 2",             "10 steps", "52%", "11 runs"],
    ["s = 4",             "15 steps", "71%",  "6 runs"],
    ["s = 6",             "17 steps", "81%",  "4 runs"],
]
bg_colors = ["f5f5f5", "e3f2fd", "e8f5e9", "f3e5f5"]
for ci, h in enumerate(hdr_data):
    cell = tbl1.rows[0].cells[ci]
    cell.text = h
    set_cell_bg(cell, "#e8eaf6")
    for run in cell.paragraphs[0].runs:
        run.bold = True; run.font.size = Pt(9.5)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for ri, row in enumerate(row_data1):
    for ci, val in enumerate(row):
        cell = tbl1.rows[ri+1].cells[ci]
        cell.text = val
        set_cell_bg(cell, "#" + bg_colors[ri])
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9.5)

fig_caption(doc, "Table 1.", "FEM computation savings per skip value (21 total time steps).")


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — Training Method
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, 3, "Training Method and Physical Parameters")
body(doc, "Each PINN window is trained with AdamW (2000 epochs). The MCO loss function:")
formula(doc,
    "\u2112  =  w_target \u00b7 \u2112_target  +  0.05 \u00b7 w_PDE \u00b7 \u2112_PDE     ,     w_i = 1/\u03c3_i\u00b2")
body(doc,
    "Adaptive weights w_i are updated via running standard deviation of each loss component, "
    "eliminating manual tuning. NAS search space: depth 2-6 layers, width 32-256 neurons, "
    "activations {Tanh, GELU, SiLU}. Optimizers: Bayesian Optimization, NSGA-II, NSGA-III.")
body(doc,
    "A356 aluminum: \u03c1=2670 kg/m\u00b3, c_p=963 J/(kg\u00b7K), k=151 W/(m\u00b7K). "
    "Boundary: convective quenching h=2000 W/(m\u00b2\u00b7K), T_\u221e=20 \u00b0C, T_0=540 \u00b0C. "
    "Governing PDE: \u03c1\u00b7c_p\u00b7\u2202T/\u2202t = \u2207\u00b7(k\u00b7\u2207T).")

# ── Page break ────────────────────────────────────────────────────────────────
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — Results: MAE Table
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, 4, "Results: Skip Operator Accuracy")
body(doc,
    "Table 2 presents MAE (deg C) for all 48 configurations. "
    "Color coding: dark green = best per row, light green < 8 C, "
    "yellow = 8-15 C, red > 15 C.")

# MAE Table
tbl2 = doc.add_table(rows=len(DOMAINS)*len(ARCHS)+1, cols=6)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = ["Domain", "Optimizer", "skip=1 (0%)", "skip=2 (52%)", "skip=4 (71%)", "skip=6 (81%)"]
for ci, h in enumerate(hdr2):
    cell = tbl2.rows[0].cells[ci]
    cell.text = h
    set_cell_bg(cell, "#c5cae9")
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cell.paragraphs[0].runs:
        run.bold = True; run.font.size = Pt(8.5)

ri = 1
for dom in DOMAINS:
    for arch in ARCHS:
        vals = [DATA[dom][arch].get(str(sk), {}).get("mae_C", None) for sk in SKIPS]
        minv = min((v for v in vals if v is not None), default=None)
        row_cells = tbl2.rows[ri].cells
        row_cells[0].text = DOM_LBL[dom]
        row_cells[1].text = ARCH_LBL[arch]
        for ci, (v, sk) in enumerate(zip(vals, SKIPS)):
            cell = row_cells[ci + 2]
            if v is None:
                cell.text = "--"; set_cell_bg(cell, "#ffffff")
            else:
                cell.text = f"{v:.2f}"
                if v == minv:   set_cell_bg(cell, "#a5d6a7")
                elif v < 8:     set_cell_bg(cell, "#e8f5e9")
                elif v < 15:    set_cell_bg(cell, "#fff9c4")
                else:           set_cell_bg(cell, "#ffcdd2")
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8.5)
        for ci in range(2):
            row_cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in row_cells[ci].paragraphs[0].runs:
                run.font.size = Pt(8.5)
        ri += 1

fig_caption(doc, "Table 2.",
    "MAE (deg C) for all 48 training runs. FEM savings in column headers. "
    "Green = good accuracy; red = MAE > 15 deg C.")

# Bar chart image
add_image(doc, "fig2_mae_per_arch.png", width_inches=6.0)
fig_caption(doc, "Figure 1.",
    "Best MAE per domain and skip value (minimum over three NAS optimizers). "
    "Each bar group = one skip value. Dashed line at 10 deg C marks a practical accuracy threshold.")

# ── Page break ────────────────────────────────────────────────────────────────
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — Visualization
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, 5, "Visualization")

add_image(doc, "fig5_skip_timeline.png", width_inches=6.2)
fig_caption(doc, "Figure 2.",
    "Skip Operator Timeline: MAE per time window (Rectangular domain, Bayesian NAS). "
    "Each line = one skip value. At skip=1 error is low and uniform. At skip=2/4 it grows "
    "moderately. At skip=6 it rises sharply in later windows.")

doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_image(doc, "fig1_thermal_fields.png", width_inches=6.2)
fig_caption(doc, "Figure 3.",
    "Predicted temperature fields on the z-midplane (Turbo colormap: blue=20 C, yellow=540 C). "
    "All four 3D domains, skip=2, Bayesian NAS, mid-simulation step. "
    "Rectangular/Cylinder: smooth surface cooling. Stacked: interface effect. "
    "L-Shape: concave-corner cooling (validated vs. numerical FD).")


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — Conclusions
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, 6, "Conclusions")

conclusions = [
    ("Best accuracy (skip=1):",
     "2.19 deg C MAE on L-Shape with NSGA-II/III. All domains achieve < 5.6 deg C."),
    ("Practical optimum (skip=2):",
     "52% FEM savings while keeping MAE < 11 deg C for all domains (recommended default)."),
    ("Fast screening (skip=6):",
     "81% FEM reduction; MAE 5.6-22.5 deg C, suitable for rapid feasibility checks."),
    ("NAS comparison:",
     "Bayesian best on simple geometries; NSGA-II/III superior on the complex L-Shape domain."),
    ("MCO weighting:",
     "Adaptive w_i = 1/sigma_i^2 eliminates manual loss tuning across all geometries."),
]
for bold_lbl, body_txt in conclusions:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, bold_lbl + "  ", bold=True, size=10, color="#1a237e")
    add_run(p, body_txt, bold=False, size=10, color="#222222")


# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Done: {OUT}")
